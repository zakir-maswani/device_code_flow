import streamlit as st
import requests
import html
import json

from msal import PublicClientApplication
from datetime import datetime, timedelta, timezone
from groq import Groq

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# CONFIG
CLIENT_ID = st.secrets.get("CLIENT_ID", "test_client")
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "test_key")

AUTHORITY = "https://login.microsoftonline.com/consumers"

SCOPES = [
    "https://graph.microsoft.com/User.Read",
    "https://graph.microsoft.com/Mail.Read"
]

app = PublicClientApplication(
    CLIENT_ID,
    authority=AUTHORITY
)

groq_client = Groq(
    api_key= GROQ_API_KEY
)

# PAGE CONFIG
st.set_page_config(
    page_title="📧 Outlook AI Report Generator",
    page_icon="⭕",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 20px;
        border-radius: 12px;
        text-align: center;
        margin: 10px 0;
    }
    .metric-value {
        font-size: 32px;
        font-weight: bold;
        margin: 10px 0;
    }
    .metric-label {
        font-size: 14px;
        opacity: 0.9;
    }
    .email-preview {
        background: #f8f9fa;
        padding: 15px;
        border-left: 4px solid #667eea;
        border-radius: 4px;
        margin: 10px 0;
    }
    .priority-critical {
        color: #ef4444;
        font-weight: bold;
    }
    .priority-high {
        color: #f59e0b;
        font-weight: bold;
    }
    .priority-medium {
        color: #10b981;
        font-weight: bold;
    }
    .priority-low {
        color: #6b7a8e;
        font-weight: bold;
    }
    .feature-item {
        padding: 6px 0;
        font-size: 13px;
        border-bottom: 1px solid #f0f0f0;
    }
</style>
""", unsafe_allow_html=True)

# SIDEBAR
with st.sidebar:
    st.markdown("## AI Report Generator")
    st.markdown("---")
    st.markdown("## Key Features")
    features = [
        ("Fetches last 7 days of emails"),
        ("AI priority classification"),
        ("Executive summary per email"),
        ("Action item extraction"),
        ("KPI dashboard overview"),
        ("Professional DOCX export"),
        ("Secure Microsoft OAuth login"),
        ("AI executive overview"),
    ]
    for text in features:
        st.markdown(f"{text}", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("## Priority Levels")
    st.markdown("🔴 **Critical** — Urgent, high-risk")
    st.markdown("🟠 **High** — Needs prompt action")
    st.markdown("🟢 **Medium** — Standard follow-up")
    st.markdown("⚪ **Low** — Informational only")
    st.markdown("---")
    st.caption("Powered by Groq · Microsoft Graph API")

st.title("Outlook AI Report Generator")

st.write(
    "Generate professional AI-powered DOCX reports from your Outlook emails with automatic formatting."
)

# ---------------------------------------------------
# COLOR SYSTEM
# ---------------------------------------------------
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
BLUE = RGBColor(0x1A, 0x56, 0xDB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT = RGBColor(0x1A, 0x20, 0x2C)
MUTED = RGBColor(0x6B, 0x7A, 0x8E)
SURFACE = RGBColor(0xF8, 0xFA, 0xFC)

RED = RGBColor(0xEF, 0x44, 0x44)
GREEN = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
SLATE = RGBColor(0x94, 0xA3, 0xB8)

RED_BG = RGBColor(0xFE, 0xF2, 0xF2)
GREEN_BG = RGBColor(0xF0, 0xFD, 0xF4)
AMBER_BG = RGBColor(0xFF, 0xFB, 0xEB)
SLATE_BG = RGBColor(0xF8, 0xFA, 0xFC)

RED_TXT = RGBColor(0x99, 0x1B, 0x1B)
GREEN_TXT = RGBColor(0x06, 0x5F, 0x46)
AMBER_TXT = RGBColor(0x92, 0x40, 0x0E)
SLATE_TXT = RGBColor(0x4A, 0x55, 0x68)

TOTAL_DXA = 9026

PRIORITY_THEME = {
    "Critical": (RED_BG, RED, RED_TXT),
    "High": (AMBER_BG, AMBER, AMBER_TXT),
    "Medium": (GREEN_BG, GREEN, GREEN_TXT),
    "Low": (SLATE_BG, SLATE, SLATE_TXT),
}

# ---------------------------------------------------
# DOCX HELPERS WITH DYNAMIC TEXT SIZING
# ---------------------------------------------------
def hex_rgb(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_rgb(rgb))
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), str(m[1]))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)


def set_col_widths(table, widths):
    table.autofit = False

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_horizontal_rule(doc, color='D1D5DB', size=6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def badge_cell(cell, text, bg, txt):
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = txt


def priority_badge(cell, priority):
    bg, _, txt = PRIORITY_THEME.get(priority, PRIORITY_THEME['Low'])
    badge_cell(cell, priority, bg, txt)


def auto_fit_text(cell, text, base_size=10):
    """Auto-fit text size based on content length"""
    p = cell.paragraphs[0]
    
    if len(text) > 200:
        font_size = Pt(8)
    elif len(text) > 100:
        font_size = Pt(9)
    else:
        font_size = Pt(base_size)
    
    p.word_wrap = True
    run = p.add_run(text)
    run.font.size = font_size
    run.font.color.rgb = TEXT
    
    return p, run


# ---------------------------------------------------
# AI ANALYSIS
# ---------------------------------------------------
def analyze_email(email: dict) -> dict:
    prompt = f"""
You are a Senior Executive Email Intelligence Analyst.

Analyze the email carefully and return ONLY ONE valid JSON object.

STRICT RULES:
- Output raw JSON only.
- Do not use markdown.
- Do not use code blocks.
- Do not provide explanations.
- Do not include any text before or after the JSON.
- Never invent facts that are not present in the email.
- If information is unavailable, use null.
- Summary must be concise and professional.
- Action item must be clear and actionable.
- If no action is required, return "No action required".

EMAIL DETAILS
-------------
Subject: {email.get("subject", "(No Subject)")}
Sender: {email.get("from", {}).get("emailAddress", {}).get("address", "Unknown")}
Preview:
{email.get("bodyPreview", "")[:1500]}

ANALYSIS INSTRUCTIONS
---------------------
1. Determine Priority:
   - Critical:
       Security incidents,
       legal escalations,
       production outages,
       executive escalations,
       financial risks,
       deadlines within 24 hours.
   - High:
       Customer issues,
       approval requests,
       urgent meetings,
       important business decisions,
       time-sensitive actions.
   - Medium:
       Standard business communications requiring attention.
   - Low:
       Informational emails,
       newsletters,
       promotions,
       automated notifications.

2. Create Executive Summary:
   - Maximum 2 sentences.
   - Mention key purpose and outcome.
   - Focus on business impact.

3. Extract Primary Action Item:
   - State the most important next step.
   - If none exists, return:
     "No action required"

4. Determine Sentiment:
   - Positive
   - Neutral
   - Negative
   - Urgent

5. Categorize Email:
   Choose ONE:
   - Meeting
   - Finance
   - Support
   - Project
   - Legal
   - HR
   - Sales
   - Security
   - Operations
   - Other

6. Detect:
   - Deadline or due date
   - Approval requests
   - Follow-up requirements

7. Assess Business Risk:
   - High
   - Medium
   - Low
   - None

OUTPUT JSON FORMAT
------------------
{{
  "priority": "Medium",
  "summary": "Professional executive summary.",
  "action_item": "Specific action or No action required",
  "sentiment": "Neutral",
  "category": "Project",
  "deadline": null,
  "approval_required": false,
  "follow_up_required": false,
  "business_risk": "Low"
}}
"""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": """
You are a strict JSON API.

Rules:
- Return ONLY valid JSON.
- Never return markdown.
- Never return code blocks.
- Never return explanations.
- Never return text outside JSON.
- Always return a single JSON object.
- Use null for unknown values.
- Follow the exact schema provided by the user.
"""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.1,
            max_tokens=400
        )

        raw = response.choices[0].message.content.strip()

        import re
        import json

        raw = re.sub(r"```(?:json)?", "", raw).strip()

        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            raw = match.group(0)

        result = json.loads(raw)

        defaults = {
            "priority": "Medium",
            "summary": "No summary available.",
            "action_item": "Review manually.",
            "sentiment": "Neutral",
            "category": "Other",
            "deadline": None,
            "approval_required": False,
            "follow_up_required": False,
            "business_risk": "None"
        }

        for key, value in defaults.items():
            result.setdefault(key, value)

        return result

    except Exception:
        return {
            "priority": "Medium",
            "summary": "AI summary could not be generated.",
            "action_item": "Review manually.",
            "sentiment": "Neutral",
            "category": "Other",
            "deadline": None,
            "approval_required": False,
            "follow_up_required": False,
            "business_risk": "None"
        }

# ---------------------------------------------------
# WEEKLY OVERVIEW
# ---------------------------------------------------
def weekly_overview(emails, analyses):

    critical = sum(
        1 for a in analyses
        if a.get("priority") == "Critical"
    )

    high = sum(
        1 for a in analyses
        if a.get("priority") == "High"
    )

    actions = [
        a.get("action_item")
        for a in analyses
        if a.get("action_item") != "No action required"
    ]

    prompt = f"""
Write a professional executive overview paragraph (max 4 sentences).

Emails analysed: {len(emails)}

Critical emails: {critical}

High priority emails: {high}

Key actions:
{' ; '.join(actions[:5]) if actions else 'None'}

Return plain professional paragraph only.
"""

    try:

        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3
        )

        return response.choices[0].message.content.strip()

    except Exception:

        return "Overview generation failed."

# ---------------------------------------------------
# PROFESSIONAL DOCX GENERATOR WITH DYNAMIC SIZING
# ---------------------------------------------------
def generate_docx(emails, analyses, overview_text):

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ---------------------------------------------------
    # COVER HEADER
    # ---------------------------------------------------
    tbl = doc.add_table(rows=3, cols=1)
    tbl.style = "Table Grid"
    set_col_widths(tbl, [Inches(7.0)])

    c0 = tbl.cell(0, 0)
    set_cell_bg(c0, NAVY)
    set_cell_margins(c0, top=100, bottom=60, left=150, right=150)
    r = c0.paragraphs[0].add_run("OUTLOOK AI REPORT")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = BLUE

    c1 = tbl.cell(1, 0)
    set_cell_bg(c1, NAVY)
    set_cell_margins(c1, top=120, bottom=120, left=150, right=150)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p1.add_run("Weekly Email Intelligence Report")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = WHITE

    c2 = tbl.cell(2, 0)
    set_cell_bg(c2, NAVY)
    set_cell_margins(c2, top=80, bottom=100, left=150, right=150)
    r = c2.paragraphs[0].add_run(
        datetime.utcnow().strftime('%d %B %Y')
    )
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE

    doc.add_paragraph()

    # ---------------------------------------------------
    # KPI SECTION
    # ---------------------------------------------------
    critical = sum(
        1 for a in analyses
        if a.get("priority") == "Critical"
    )

    high = sum(
        1 for a in analyses
        if a.get("priority") == "High"
    )

    actions = sum(
        1 for a in analyses
        if a.get("action_item") != "No action required"
    )

    kpi = doc.add_table(rows=2, cols=4)
    kpi.style = "Table Grid"
    set_col_widths(kpi, [Inches(1.75), Inches(1.75), Inches(1.75), Inches(1.75)])

    values = [
        (str(len(emails)), "Total Emails", BLUE),
        (str(critical), "Critical", RED),
        (str(high), "High Priority", AMBER),
        (str(actions), "Action Items", GREEN)
    ]

    for idx, (num, label, color) in enumerate(values):

        cell1 = kpi.cell(0, idx)
        cell2 = kpi.cell(1, idx)

        set_cell_bg(cell1, WHITE)
        set_cell_bg(cell2, SURFACE)
        set_cell_margins(cell1, top=80, bottom=60, left=80, right=80)
        set_cell_margins(cell2, top=60, bottom=80, left=80, right=80)

        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p1.add_run(num)
        r1.bold = True
        r1.font.size = Pt(24)
        r1.font.color.rgb = color

        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED

    doc.add_paragraph()

    # ---------------------------------------------------
    # OVERVIEW
    # ---------------------------------------------------
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    rh = h.add_run("EXECUTIVE OVERVIEW")
    rh.bold = True
    rh.font.size = Pt(13)
    rh.font.color.rgb = NAVY

    add_horizontal_rule(doc, color='1A56DB', size=8)

    ov = doc.add_table(rows=1, cols=2)
    ov.style = "Table Grid"
    set_col_widths(ov, [Inches(0.15), Inches(6.85)])

    left = ov.cell(0, 0)
    right = ov.cell(0, 1)

    set_cell_bg(left, BLUE)
    set_cell_bg(right, WHITE)
    set_cell_margins(right, top=100, bottom=100, left=120, right=120)

    left.width = Inches(0.1)

    pr = right.paragraphs[0]
    pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = pr.add_run(overview_text)
    rr.font.size = Pt(11)
    rr.font.color.rgb = TEXT

    doc.add_paragraph()

    # ---------------------------------------------------
    # EMAIL ANALYSIS
    # ---------------------------------------------------
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    rh = heading.add_run("EMAIL ANALYSIS")
    rh.bold = True
    rh.font.size = Pt(13)
    rh.font.color.rgb = NAVY

    add_horizontal_rule(doc, color='1A56DB', size=8)

    for idx, (email, analysis) in enumerate(
        zip(emails, analyses),
        start=1
    ):

        subject = email.get("subject", "No Subject")[:100]

        sender = (
            email.get("from", {})
            .get("emailAddress", {})
            .get("address", "Unknown")
        )

        received = email.get(
            "receivedDateTime",
            ""
        )[:10]

        preview = email.get(
            "bodyPreview",
            ""
        )[:250]

        priority = analysis.get("priority", "Medium")

        card = doc.add_table(rows=1, cols=3)
        card.style = "Table Grid"
        set_col_widths(card, [Inches(0.6), Inches(5.4), Inches(1.0)])

        c1 = card.cell(0, 0)
        c2 = card.cell(0, 1)
        c3 = card.cell(0, 2)

        set_cell_bg(c1, SURFACE)
        set_cell_bg(c2, SURFACE)
        set_cell_margins(c1, top=80, bottom=80, left=80, right=80)
        set_cell_margins(c2, top=80, bottom=80, left=100, right=100)

        r = c1.paragraphs[0].add_run(f"#{idx}")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = MUTED

        p2 = c2.paragraphs[0]
        p2.word_wrap = True
        r = p2.add_run(subject)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT

        priority_badge(c3, priority)

        meta = doc.add_table(rows=2, cols=4)
        meta.style = "Table Grid"
        set_col_widths(meta, [Inches(1.8), Inches(1.8), Inches(1.6), Inches(1.6)])

        labels = [
            "FROM",
            "DATE",
            "CATEGORY",
            "SENTIMENT"
        ]

        values = [
            sender[:40],
            received,
            analysis.get("category", ""),
            analysis.get("sentiment", "")
        ]

        for i in range(4):

            top = meta.cell(0, i)
            bottom = meta.cell(1, i)

            set_cell_bg(top, SURFACE)
            set_cell_bg(bottom, WHITE)
            set_cell_margins(top, top=60, bottom=40, left=80, right=80)
            set_cell_margins(bottom, top=60, bottom=80, left=80, right=80)

            rt = top.paragraphs[0].add_run(labels[i])
            rt.bold = True
            rt.font.size = Pt(8)
            rt.font.color.rgb = MUTED

            rb = bottom.paragraphs[0]
            rb.word_wrap = True
            r = rb.add_run(str(values[i]))
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT

        body = doc.add_table(rows=3, cols=2)
        body.style = "Table Grid"
        set_col_widths(body, [Inches(1.7), Inches(5.2)])

        rows = [
            (
                "AI SUMMARY",
                analysis.get("summary", "")[:180]
            ),
            (
                "ACTION REQUIRED",
                analysis.get("action_item", "")[:120]
            ),
            (
                "EMAIL PREVIEW",
                preview
            )
        ]

        for r_idx, (label, value) in enumerate(rows):

            lc = body.cell(r_idx, 0)
            vc = body.cell(r_idx, 1)

            set_cell_bg(lc, SURFACE)
            set_cell_bg(vc, WHITE)
            set_cell_margins(lc, top=70, bottom=70, left=100, right=80)
            set_cell_margins(vc, top=70, bottom=70, left=100, right=100)

            rl = lc.paragraphs[0].add_run(label)
            rl.bold = True
            rl.font.size = Pt(8)
            rl.font.color.rgb = MUTED

            rv = vc.paragraphs[0]
            rv.word_wrap = True
            rv.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = rv.add_run(str(value))
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT

        doc.add_paragraph()

    # ---------------------------------------------------
    # ACTION ITEMS PAGE
    # ---------------------------------------------------
    doc.add_page_break()

    action_heading = doc.add_paragraph()
    action_heading.paragraph_format.space_before = Pt(12)
    action_heading.paragraph_format.space_after = Pt(6)
    ra = action_heading.add_run("ACTION ITEMS SUMMARY")
    ra.bold = True
    ra.font.size = Pt(14)
    ra.font.color.rgb = NAVY

    add_horizontal_rule(doc, color='1A56DB', size=8)

    actions_table = doc.add_table(rows=1, cols=4)
    actions_table.style = "Table Grid"
    set_col_widths(actions_table, [Inches(2.6), Inches(1.0), Inches(1.2), Inches(2.2)])

    headers = [
        "Subject",
        "Priority",
        "Category",
        "Action"
    ]

    for idx, h in enumerate(headers):

        cell = actions_table.cell(0, idx)

        set_cell_bg(cell, NAVY)
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE

    for email, analysis in zip(emails, analyses):

        if analysis.get("action_item") == "No action required":
            continue

        row = actions_table.add_row().cells
        
        for cell in row:
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)

        p0 = row[0].paragraphs[0]
        p0.word_wrap = True
        p0.add_run(email.get("subject", "")[:50])
        p0.runs[0].font.size = Pt(9)

        p1 = row[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(analysis.get("priority", ""))
        p1.runs[0].font.size = Pt(9)

        p2 = row[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(analysis.get("category", ""))
        p2.runs[0].font.size = Pt(9)

        p3 = row[3].paragraphs[0]
        p3.word_wrap = True
        p3.add_run(analysis.get("action_item", "")[:60])
        p3.runs[0].font.size = Pt(9)

    output_file = "weekly_ai_report.docx"

    doc.save(output_file)

    return output_file

# ---------------------------------------------------
# LOGIN
# ---------------------------------------------------
if "access_token" not in st.session_state:

    st.info(
        "🔓 Login with your Microsoft account to continue"
    )

    if st.button("🔐 Login with Microsoft", use_container_width=True):

        try:

            flow = app.initiate_device_flow(
                scopes=SCOPES
            )

            if "user_code" not in flow:

                st.error("❌ Device flow failed")

                st.stop()

            st.markdown("### 👇 Complete Login")

            st.code(flow["user_code"])

            st.markdown(
                f"""
Visit: **{flow['verification_uri']}**

Enter the code above.
"""
            )

            with st.spinner(
                "Waiting for login..."
            ):

                result = app.acquire_token_by_device_flow(
                    flow
                )

            if "access_token" in result:

                st.session_state["access_token"] = (
                    result["access_token"]
                )

                st.success(
                    "✅ Login successful"
                )

                st.rerun()

            else:

                st.error(
                    result.get(
                        "error_description",
                        "Login failed"
                    )
                )

        except Exception as e:

            st.error(f"❌ {str(e)}")

# ---------------------------------------------------
# AFTER LOGIN
# ---------------------------------------------------
else:

    access_token = st.session_state["access_token"]

    col1, col2 = st.columns([3, 1])

    with col1:
        generate_btn = st.button(
            "📄 Generate AI Report",
            use_container_width=True
        )

    with col2:
        logout_btn = st.button(
            "🚪 Logout",
            use_container_width=True
        )

    if logout_btn:

        st.session_state.clear()

        st.rerun()

    if generate_btn:

        try:

            with st.spinner(
                "📬 Fetching emails..."
            ):

                last_week = (
                    datetime.now(timezone.utc)
                    - timedelta(days=7)
                ).strftime("%Y-%m-%dT%H:%M:%SZ")

                url = (
                    "https://graph.microsoft.com/v1.0/me/messages"
                    f"?$filter=receivedDateTime ge {last_week}"
                    "&$top=20"
                    "&$orderby=receivedDateTime DESC"
                )

                headers = {
                    "Authorization":
                    f"Bearer {access_token}"
                }

                response = requests.get(
                    url,
                    headers=headers,
                    timeout=20
                )

                if response.status_code == 401:

                    st.error(
                        "❌ Session expired. Please login again."
                    )

                    st.session_state.clear()

                    st.stop()

                if response.status_code != 200:

                    st.error(
                        f"❌ Graph API Error: "
                        f"{response.status_code}"
                    )

                    try:
                        st.json(response.json())
                    except Exception:
                        st.write(response.text)

                    st.stop()

                emails = response.json().get(
                    "value",
                    []
                )

            st.success(
                f"{len(emails)} emails fetched"
            )

            if not emails:

                st.info(
                    "No emails found in last 7 days"
                )

                st.stop()

            analyses = []

            progress = st.progress(0)

            with st.spinner(
                "Analysing emails using AI..."
            ):

                total = len(emails)

                for idx, email in enumerate(emails):

                    result = analyze_email(email)

                    analyses.append(result)

                    progress.progress(
                        (idx + 1) / total
                    )

            with st.spinner(
                "Generating executive overview..."
            ):

                overview = weekly_overview(
                    emails,
                    analyses
                )

            with st.spinner(
                "Generating professional DOCX report..."
            ):

                report_file = generate_docx(
                    emails,
                    analyses,
                    overview
                )

            st.success(
                "AI Report Generated Successfully"
            )

            # ---------------------------------------------------
            # LIVE PREVIEW SECTION
            # ---------------------------------------------------
            st.subheader("Report Preview & Statistics")

            kpi_col1, kpi_col2, kpi_col3, kpi_col4 = st.columns(4)

            with kpi_col1:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Total Emails</div>
                    <div class="metric-value">{len(emails)}</div>
                </div>
                """, unsafe_allow_html=True)

            with kpi_col2:
                critical_count = sum(1 for a in analyses if a.get("priority") == "Critical")
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Critical</div>
                    <div class="metric-value" style="color: #ef4444;">{critical_count}</div>
                </div>
                """, unsafe_allow_html=True)

            with kpi_col3:
                high_count = sum(1 for a in analyses if a.get("priority") == "High")
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">High Priority</div>
                    <div class="metric-value" style="color: #f59e0b;">{high_count}</div>
                </div>
                """, unsafe_allow_html=True)

            with kpi_col4:
                action_count = sum(1 for a in analyses if a.get("action_item") != "No action required")
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-label">Action Items</div>
                    <div class="metric-value" style="color: #10b981;">{action_count}</div>
                </div>
                """, unsafe_allow_html=True)

            st.divider()

            # Overview
            st.subheader("Executive Overview")
            st.info(overview)

            st.divider()

            # Download Button
            with open(report_file, "rb") as file:

                st.download_button(
                    label="Download DOCX Report",
                    data=file,
                    file_name=report_file,
                    mime=(
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    ),
                    use_container_width=True
                )

        except requests.exceptions.Timeout:

            st.error("⏱️ Request timeout")

        except requests.exceptions.ConnectionError:

            st.error("🌐 Network connection error")

        except Exception as e:

            st.error(f"❌ {str(e)}")
